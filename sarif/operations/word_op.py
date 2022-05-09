"""
Generate summary of given JSON or given JSON Document in to Microsoft Word Documents.
This functionality uses a python module called

python-docx - a Python library for creating and updating Microsoft Word (.docx) files.

https://python-docx.readthedocs.io/

"""

from datetime import datetime
import os

import docx
from docx import oxml
from docx import shared
from docx.enum import text
from docx.oxml import ns

from sarif import charts, sarif_file
from sarif.sarif_file import SarifFileSet


def generate_word_docs_from_sarif_inputs(
    input_files: SarifFileSet, image_file: str, output: str, output_multiple_files: bool
):
    """
    Convert SARIF input to Word file output.
    """
    if not input_files:
        raise ValueError("No input files specified!")

    output_file = output
    output_file_name = output
    if output_multiple_files:
        for input_file in input_files:
            output_file_name = input_file.get_file_name_without_extension() + ".docx"
            print(
                "Writing Word summary of",
                input_file.get_file_name(),
                "to",
                output_file_name,
            )
            _generate_word_summary(
                input_file,
                os.path.join(output, output_file_name),
                image_file,
            )
        output_file_name = "static_analysis_output.docx"
        output_file = os.path.join(output, output_file_name)

    source_description = input_files.get_description()
    print("Writing Word summary of", source_description, "to", output_file_name)
    _generate_word_summary(input_files, output_file, image_file)


def _generate_word_summary(sarif_data, output_file, image_file):

    # Create a new document
    document = docx.Document()

    _add_heading_and_highlevel_info(document, sarif_data, output_file, image_file)
    _dump_errors_summary_by_sev(document, sarif_data)
    _dump_each_error_in_detail(document, sarif_data)

    # finally, save the document.
    document.save(output_file)


def _add_heading_and_highlevel_info(document, sarif_data, output_file, image_path):
    tool_name = ", ".join(sarif_data.get_distinct_tool_names())
    heading = f"Sarif Summary: {tool_name}"

    if image_path:
        document.add_picture(image_path)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_heading(heading, 0)
    document.add_paragraph(f"Document generated on: {datetime.now()}")

    sevs = ", ".join(sarif_file.SARIF_SEVERITIES)
    document.add_paragraph(
        f"Total number of various severities ({sevs}): {sarif_data.get_result_count()}"
    )
    filter_stats = sarif_data.get_filter_stats()
    if filter_stats:
        document.add_paragraph(f"Results were filtered by {filter_stats}.")

    pie_chart_image_file_path = output_file.replace(".docx", "_severity_pie_chart.png")
    if charts.generate_severity_pie_chart(sarif_data, pie_chart_image_file_path):
        document.add_picture(pie_chart_image_file_path)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()


def _dump_errors_summary_by_sev(document, sarif_data):
    """
    For each severity level (in priority order): create a list of the errors of
    that severity, print out how many there are and then do some further analysis
    of which error codes are present.
    """

    severities = sarif_file.SARIF_SEVERITIES
    sev_to_records = sarif_data.get_records_grouped_by_severity()
    for severity in severities:
        errors_of_severity = sev_to_records.get(severity, [])
        document.add_heading(
            f"Severity : {severity} [ {len(errors_of_severity)} ]", level=1
        )

        # Go through the list of errors and create a dictionary of each error code
        # present to how many times that error code occurs. Sort this dict and print
        # out in descending order.
        dict_of_error_codes = {}
        for error in errors_of_severity:
            issue_code = error["Code"]
            dict_of_error_codes[issue_code] = dict_of_error_codes.get(issue_code, 0) + 1
        sorted_dict = sorted(
            dict_of_error_codes.items(), key=lambda x: x[1], reverse=True
        )
        if sorted_dict:
            for error in sorted_dict:
                document.add_paragraph(f"{error[0]}: {error[1]}", style="List Bullet")
        else:
            document.add_paragraph("None", style="List Bullet")


def _dump_each_error_in_detail(document, sarif_data):
    """
    Write out the errors to a table so that a human can do further analysis.
    """
    document.add_page_break()

    severities = sarif_file.SARIF_SEVERITIES
    sev_to_records = sarif_data.get_records_grouped_by_severity()
    for severity in severities:
        errors_of_severity = sev_to_records.get(severity, [])
        sorted_errors_by_severity = sorted(errors_of_severity, key=lambda x: x["Code"])
        # Sample:
        # [{'Location': 'C:\\Max\\AccessionAndroid\\scripts\\parse_coverage.py', 'Line': 119,
        #       'Severity': 'error', 'Code': 'DS126186 Disabled certificate validation'},
        # {'Location': 'C:\\Max\\AccessionAndroid\\scripts\\parse_code_stats.py', 'Line': 61,
        #       'Severity': 'error', 'Code': 'DS126186 Disabled certificate validation'},
        # ]
        if errors_of_severity:
            document.add_heading(f"Severity : {severity}", level=2)
            table = document.add_table(rows=1 + len(errors_of_severity), cols=3)

            table.style = "Table Grid"  # ColorfulGrid-Accent5'
            table.autofit = False

            table.alignment = text.WD_TAB_ALIGNMENT.CENTER

            # Cell widths
            widths = [shared.Inches(2), shared.Inches(4), shared.Inches(0.5)]

            # To avoid performance problems with large tables, prepare the entries first in this
            # list, then iterate the table cells and copy them in.
            # First populate the header row
            cells_text = ["Code", "Location", "Line"]

            hdr_cells = table.rows[0].cells
            for i in range(3):
                table.rows[0].cells[i]._tc.get_or_add_tcPr().append(
                    oxml.parse_xml(
                        r'<w:shd {} w:fill="5fe3d8"/>'.format(ns.nsdecls("w"))
                    )
                )
                run = hdr_cells[i].paragraphs[0].add_run(cells_text[i])
                run.bold = True
                hdr_cells[i].paragraphs[
                    0
                ].alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER
                hdr_cells[i].width = widths[i]

            for eachrow in sorted_errors_by_severity:
                cells_text += [
                    eachrow["Code"],
                    eachrow["Location"],
                    str(eachrow["Line"]),
                ]

            # Note: using private property table._cells to avoid performance issue.  See
            # https://stackoverflow.com/a/69105798/316578
            col_index = 0
            for (cell, cell_text) in zip(table._cells, cells_text):
                cell.text = cell_text
                cell.width = widths[col_index]
                col_index = col_index + 1 if col_index < 2 else 0
        else:
            document.add_heading(f"Severity : {severity}", level=2)
            document.add_paragraph("None", style="List Bullet")
